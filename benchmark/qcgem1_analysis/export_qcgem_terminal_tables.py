#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Export QC-GEM W8A16/W4A16 benchmark data to CSV, Markdown and Word documents.
"""

from __future__ import annotations

import csv
import os
from pathlib import Path

ZH_ROOT = Path(__file__).resolve().parent

# W8A16 benchmark data
W8A16_DATA = [
    (1, "Down projection", "32768x3584", 32768, 3584, 1024, 2.560480, 2.564400, 93.791, 0.998),
    (2, "Down projection", "16384x3584", 16384, 3584, 1024, 1.400416, 1.391984, 86.394, 1.006),
    (3, "Down projection", "8192x3584", 8192, 3584, 1024, 0.815424, 0.811264, 74.118, 1.005),
    (4, "Down projection", "4096x3584", 4096, 3584, 1024, 0.510176, 0.502720, 59.804, 1.015),
    (5, "Down projection", "2048x3584", 2048, 3584, 1024, 0.354688, 0.348608, 43.121, 1.017),
    (6, "Down projection", "1024x3584", 1024, 3584, 1024, 0.277248, 0.313424, 23.981, 0.885),
    (7, "Down projection", "512x3584", 512, 3584, 1024, 0.235104, 0.279296, 13.456, 0.842),
    (8, "Down projection", "256x3584", 256, 3584, 1024, 0.219392, 0.261328, 7.190, 0.840),
    (9, "Up projection", "32768x1024", 32768, 1024, 3584, 2.637664, 2.619440, 91.820, 1.007),
    (10, "Up projection", "16384x1024", 16384, 1024, 3584, 1.406016, 1.396032, 86.144, 1.007),
    (11, "Up projection", "8192x1024", 8192, 1024, 3584, 0.797296, 0.789568, 76.155, 1.010),
    (12, "Up projection", "4096x1024", 4096, 1024, 3584, 0.490720, 0.483008, 62.245, 1.016),
    (13, "Up projection", "2048x1024", 2048, 1024, 3584, 0.345712, 0.338912, 44.355, 1.020),
    (14, "Gate projection", "32768x1024", 32768, 1024, 7168, 5.226288, 5.198688, 92.530, 1.005),
    (15, "Gate projection", "16384x1024", 16384, 1024, 7168, 2.754464, 2.754240, 87.327, 1.000),
    (16, "Gate projection", "8192x1024", 8192, 1024, 7168, 1.548672, 1.540960, 78.042, 1.005),
    (17, "Gate projection", "4096x1024", 4096, 1024, 7168, 0.943376, 0.934080, 64.373, 1.010),
    (18, "Gate projection", "2048x1024", 2048, 1024, 7168, 0.636144, 0.630864, 47.657, 1.008),
    (19, "Router", "32768x1024", 32768, 1024, 128, 0.186880, 0.303936, 28.262, 0.615),
    (20, "Router", "16384x1024", 16384, 1024, 128, 0.187488, 0.274464, 15.649, 0.683),
    (21, "Router", "4096x1024", 4096, 1024, 128, 0.182944, 0.238272, 4.506, 0.768),
]

# W4A16 benchmark data
W4A16_DATA = [
    (1, "Down projection", "32768x3584", 32768, 3584, 1024, 2.543360, 2.541600, 94.633, 1.001),
    (2, "Down projection", "16384x3584", 16384, 3584, 1024, 1.358880, 1.348704, 89.166, 1.008),
    (3, "Down projection", "8192x3584", 8192, 3584, 1024, 0.762928, 0.756416, 79.493, 1.009),
    (4, "Down projection", "4096x3584", 4096, 3584, 1024, 0.450848, 0.443664, 67.765, 1.016),
    (5, "Down projection", "2048x3584", 2048, 3584, 1024, 0.282528, 0.372512, 40.354, 0.758),
    (6, "Down projection", "1024x3584", 1024, 3584, 1024, 0.207328, 0.363968, 20.651, 0.570),
    (7, "Down projection", "512x3584", 512, 3584, 1024, 0.203360, 0.326848, 11.498, 0.622),
    (8, "Down projection", "256x3584", 256, 3584, 1024, 0.204224, 0.303104, 6.199, 0.674),
    (9, "Up projection", "32768x1024", 32768, 1024, 3584, 2.617888, 2.618096, 91.868, 1.000),
    (10, "Up projection", "16384x1024", 16384, 1024, 3584, 1.366160, 1.352608, 88.909, 1.010),
    (11, "Up projection", "8192x1024", 8192, 1024, 3584, 0.742912, 0.734800, 81.831, 1.011),
    (12, "Up projection", "4096x1024", 4096, 1024, 3584, 0.430704, 0.423664, 70.964, 1.017),
    (13, "Up projection", "2048x1024", 2048, 1024, 3584, 0.274272, 0.362304, 41.491, 0.757),
    (14, "Gate projection", "32768x1024", 32768, 1024, 7168, 5.192416, 5.198560, 92.533, 0.999),
    (15, "Gate projection", "16384x1024", 16384, 1024, 7168, 2.674864, 2.674784, 89.921, 1.000),
    (16, "Gate projection", "8192x1024", 8192, 1024, 7168, 1.443232, 1.437728, 83.645, 1.004),
    (17, "Gate projection", "4096x1024", 4096, 1024, 7168, 0.819456, 0.813696, 73.897, 1.007),
    (18, "Gate projection", "2048x1024", 2048, 1024, 7168, 0.506368, 0.501792, 59.915, 1.009),
    (19, "Router", "32768x1024", 32768, 1024, 128, 0.201216, 0.344624, 24.926, 0.584),
    (20, "Router", "16384x1024", 16384, 1024, 128, 0.201840, 0.301728, 14.235, 0.669),
    (21, "Router", "4096x1024", 4096, 1024, 128, 0.200960, 0.273616, 3.924, 0.734),
]


def save_csv(data, filename):
    fieldnames = ["idx", "category", "shape", "m", "n", "k", "torch_ms", "gems_ms", "tflops", "speedup"]
    with open(filename, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        for row in data:
            writer.writerow({
                "idx": row[0],
                "category": row[1],
                "shape": row[2],
                "m": row[3],
                "n": row[4],
                "k": row[5],
                "torch_ms": row[6],
                "gems_ms": row[7],
                "tflops": row[8],
                "speedup": row[9],
            })


def write_md(rows_w8, rows_w4, out):
    def calc_stats(data):
        speedups = [r[9] for r in data]
        tflops = [r[8] for r in data]
        return {
            "mean_sp": sum(speedups) / len(speedups),
            "max_sp": max(speedups),
            "min_sp": min(speedups),
            "mean_tf": sum(tflops) / len(tflops),
            "max_tf": max(tflops),
            "count_ge1": sum(1 for s in speedups if s >= 1.0),
            "total": len(speedups),
        }

    w8_stats = calc_stats(rows_w8)
    w4_stats = calc_stats(rows_w4)

    lines = [
        "# QC-GEM W8A16/W4A16 vs PyTorch FP16 GEMM — Benchmark Summary (H20 GPU)",
        "",
        "**环境**: NVIDIA H20 (sm_90), PyTorch, Triton, FlagGems.",
        "",
        "对比项: **QC-GEM** (量化内核) 相对 **PyTorch FP16 GEMM** 的延迟与算力。",
        "",
        "---",
        "",
        "## W8A16 (INT8 权重 × FP16 激活)",
        "",
        "| # | Category | Shape | M | N | K | Torch (ms) | Gems (ms) | TFLOPS | Speedup |",
        "|---:|---|---|---:|---:|---:|---:|---:|---:|---:|",
    ]

    for r in rows_w8:
        lines.append(
            f"| {r[0]} | {r[1]} | {r[2]} | {r[3]} | {r[4]} | {r[5]} | {r[6]:.4f} | {r[7]:.4f} | {r[8]:.2f} | {r[9]:.3f}× |"
        )

    lines.extend([
        "",
        f"**统计**: 均加速比 {w8_stats['mean_sp']:.4f}×, 最大 {w8_stats['max_sp']:.3f}×, 最小 {w8_stats['min_sp']:.3f}×, "
        f"≥1.0配置: {w8_stats['count_ge1']}/{w8_stats['total']}, 均TFLOPS: {w8_stats['mean_tf']:.2f}, "
        f"最大TFLOPS: {w8_stats['max_tf']:.2f}",
        "",
        "---",
        "",
        "## W4A16 (INT4 权重 × FP16 激活)",
        "",
        "| # | Category | Shape | M | N | K | Torch (ms) | Gems (ms) | TFLOPS | Speedup |",
        "|---:|---|---|---:|---:|---:|---:|---:|---:|---:|",
    ])

    for r in rows_w4:
        lines.append(
            f"| {r[0]} | {r[1]} | {r[2]} | {r[3]} | {r[4]} | {r[5]} | {r[6]:.4f} | {r[7]:.4f} | {r[8]:.2f} | {r[9]:.3f}× |"
        )

    lines.extend([
        "",
        f"**统计**: 均加速比 {w4_stats['mean_sp']:.4f}×, 最大 {w4_stats['max_sp']:.3f}×, 最小 {w4_stats['min_sp']:.3f}×, "
        f"≥1.0配置: {w4_stats['count_ge1']}/{w4_stats['total']}, 均TFLOPS: {w4_stats['mean_tf']:.2f}, "
        f"最大TFLOPS: {w4_stats['max_tf']:.2f}",
        "",
        "---",
        "",
        "## 性能对比总结",
        "",
        "| 指标 | W8A16 | W4A16 |",
        "|---|---:|---:|",
        f"| 均加速比 | {w8_stats['mean_sp']:.4f}× | {w4_stats['mean_sp']:.4f}× |",
        f"| 最大加速比 | {w8_stats['max_sp']:.3f}× | {w4_stats['max_sp']:.3f}× |",
        f"| 最小加速比 | {w8_stats['min_sp']:.3f}× | {w4_stats['min_sp']:.3f}× |",
        f"| ≥1.0配置占比 | {w8_stats['count_ge1']}/{w8_stats['total']} | {w4_stats['count_ge1']}/{w4_stats['total']} |",
        f"| 均TFLOPS | {w8_stats['mean_tf']:.2f} | {w4_stats['mean_tf']:.2f} |",
        f"| 最大TFLOPS | {w8_stats['max_tf']:.2f} | {w4_stats['max_tf']:.2f} |",
    ])

    out.write_text("\n".join(lines), encoding="utf-8")


def write_docx(rows_w8, rows_w4, out):
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except ImportError:
        print("[WARN] python-docx not installed. Skipping DOCX generation.")
        return

    def calc_stats(data):
        speedups = [r[9] for r in data]
        tflops = [r[8] for r in data]
        return {
            "mean_sp": sum(speedups) / len(speedups),
            "max_sp": max(speedups),
            "min_sp": min(speedups),
            "mean_tf": sum(tflops) / len(tflops),
            "max_tf": max(tflops),
            "count_ge1": sum(1 for s in speedups if s >= 1.0),
            "total": len(speedups),
        }

    doc = Document()
    h = doc.add_heading("QC-GEM W8A16/W4A16 vs PyTorch FP16 GEMM — Benchmark Summary (H20 GPU)", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("环境: NVIDIA H20 (sm_90), PyTorch, Triton, FlagGems.")
    doc.add_paragraph("对比项: QC-GEM (量化内核) 相对 PyTorch FP16 GEMM 的延迟与算力。")

    cols = ["#", "Category", "Shape", "M", "N", "K", "Torch (ms)", "Gems (ms)", "TFLOPS", "Speedup"]

    doc.add_heading("W8A16 (INT8 权重 × FP16 激活)", level=2)
    t1 = doc.add_table(rows=1 + len(rows_w8), cols=len(cols))
    t1.style = "Table Grid"
    for j, c in enumerate(cols):
        t1.rows[0].cells[j].text = c
        for run in t1.rows[0].cells[j].paragraphs[0].runs:
            run.bold = True
    for i, r in enumerate(rows_w8, start=1):
        row = t1.rows[i].cells
        vals = [str(r[0]), r[1], r[2], str(r[3]), str(r[4]), str(r[5]),
                f"{r[6]:.4f}", f"{r[7]:.4f}", f"{r[8]:.2f}", f"{r[9]:.3f}×"]
        for j, v in enumerate(vals):
            row[j].text = v

    w8_stats = calc_stats(rows_w8)
    doc.add_paragraph(
        f"统计: 均加速比 {w8_stats['mean_sp']:.4f}×, 最大 {w8_stats['max_sp']:.3f}×, "
        f"最小 {w8_stats['min_sp']:.3f}×, ≥1.0配置: {w8_stats['count_ge1']}/{w8_stats['total']}, "
        f"均TFLOPS: {w8_stats['mean_tf']:.2f}, 最大TFLOPS: {w8_stats['max_tf']:.2f}"
    )

    doc.add_heading("W4A16 (INT4 权重 × FP16 激活)", level=2)
    t2 = doc.add_table(rows=1 + len(rows_w4), cols=len(cols))
    t2.style = "Table Grid"
    for j, c in enumerate(cols):
        t2.rows[0].cells[j].text = c
        for run in t2.rows[0].cells[j].paragraphs[0].runs:
            run.bold = True
    for i, r in enumerate(rows_w4, start=1):
        row = t2.rows[i].cells
        vals = [str(r[0]), r[1], r[2], str(r[3]), str(r[4]), str(r[5]),
                f"{r[6]:.4f}", f"{r[7]:.4f}", f"{r[8]:.2f}", f"{r[9]:.3f}×"]
        for j, v in enumerate(vals):
            row[j].text = v

    w4_stats = calc_stats(rows_w4)
    doc.add_paragraph(
        f"统计: 均加速比 {w4_stats['mean_sp']:.4f}×, 最大 {w4_stats['max_sp']:.3f}×, "
        f"最小 {w4_stats['min_sp']:.3f}×, ≥1.0配置: {w4_stats['count_ge1']}/{w4_stats['total']}, "
        f"均TFLOPS: {w4_stats['mean_tf']:.2f}, 最大TFLOPS: {w4_stats['max_tf']:.2f}"
    )

    doc.add_heading("性能对比总结", level=2)
    doc.add_paragraph(f"W8A16 均加速比: {w8_stats['mean_sp']:.4f}× | W4A16 均加速比: {w4_stats['mean_sp']:.4f}×")
    doc.add_paragraph(f"W8A16 最大加速比: {w8_stats['max_sp']:.3f}× | W4A16 最大加速比: {w4_stats['max_sp']:.3f}×")
    doc.add_paragraph(f"W8A16 ≥1.0配置: {w8_stats['count_ge1']}/{w8_stats['total']} | W4A16 ≥1.0配置: {w4_stats['count_ge1']}/{w4_stats['total']}")
    doc.add_paragraph(f"W8A16 最大TFLOPS: {w8_stats['max_tf']:.2f} | W4A16 最大TFLOPS: {w4_stats['max_tf']:.2f}")

    doc.save(str(out))


def main():
    csv_w8 = ZH_ROOT / "qcgem1_w8a16_data.csv"
    csv_w4 = ZH_ROOT / "qcgem1_w4a16_data.csv"
    md_out = ZH_ROOT / "qcgem1_benchmark_summary.md"
    docx_out = ZH_ROOT / "qcgem1_benchmark_summary.docx"

    save_csv(W8A16_DATA, csv_w8)
    print(f"[OK] CSV W8A16 → {csv_w8}")
    save_csv(W4A16_DATA, csv_w4)
    print(f"[OK] CSV W4A16 → {csv_w4}")

    write_md(W8A16_DATA, W4A16_DATA, md_out)
    print(f"[OK] MD → {md_out}")

    write_docx(W8A16_DATA, W4A16_DATA, docx_out)
    print(f"[OK] DOCX → {docx_out}")


if __name__ == "__main__":
    main()
