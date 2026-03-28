#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Read qcmoe_w8a16_data.csv / qcmoe_w4a16_data.csv from FlagGems benchmark,
write editable Markdown + Word summary (H20 QC-MoE vs vLLM-style FP16).
"""

from __future__ import annotations

import csv
import os
from pathlib import Path

# Default CSV locations (override with env FLAGGEMS_ROOT)
ZH_ROOT = Path(__file__).resolve().parent
# CSV 优先与本脚本同目录；否则用 FlagGems 下 qcmoe_analysis
DEFAULT_FLAGGEMS = ZH_ROOT.parent / "FlagEnv" / "FlagGems" / "benchmark" / "qcmoe_analysis"


def load_csv(path: Path) -> list[dict]:
    with open(path, encoding="utf-8", newline="") as f:
        return list(csv.DictReader(f))


def fmt_err(v: str) -> str:
    v = (v or "").strip().lower()
    if v in ("nan", "", "none"):
        return "—"
    return v


def write_md(rows_w8: list[dict], rows_w4: list[dict], out: Path) -> None:
    lines = [
        "# QC-MoE vs vLLM-style FP16 MoE — Benchmark Summary (H20 GPU)",
        "",
        "**环境（示例）**：NVIDIA H20 (sm_90)，PyTorch 2.10.0+cu128，Triton 3.6.0。",
        "",
        "对比项：**QC-MoE**（量化内核）相对 **vLLM 风格 FP16 SwiGLU MoE**（纯 PyTorch 参考）的延迟与算力。",
        "",
        "---",
        "",
        "## W8A16（INT8 权重 × FP16 激活）",
        "",
        "| # | Category | Shape (S,H,K) | QC W8A16 (ms) | FP16 ref (ms) | QC TFLOPS | FP16 TFLOPS | Speedup | MaxAbsErr |",
        "|---:|---|---|---:|---:|---:|---:|---:|---|",
    ]
    for r in rows_w8:
        lines.append(
            "| {idx} | {cat} | {shape} | {qc} | {fp} | {qct} | {fpt} | {sp} | {err} |".format(
                idx=r["idx"],
                cat=r["category"],
                shape=r["shape"],
                qc=r["qc_ms"],
                fp=r["fp16_ms"],
                qct=r["qc_tflops"],
                fpt=r["fp16_tflops"],
                sp=f"{float(r['speedup']):.2f}×",
                err=fmt_err(r.get("max_abs_err", "")),
            )
        )
    lines.extend(
        [
            "",
            "## W4A16（INT4 权重 × FP16 激活）",
            "",
            "| # | Category | Shape (S,H,K) | QC W4A16 (ms) | FP16 ref (ms) | QC TFLOPS | FP16 TFLOPS | Speedup | MaxAbsErr |",
            "|---:|---|---|---:|---:|---:|---:|---:|---|",
        ]
    )
    for r in rows_w4:
        lines.append(
            "| {idx} | {cat} | {shape} | {qc} | {fp} | {qct} | {fpt} | {sp} | {err} |".format(
                idx=r["idx"],
                cat=r["category"],
                shape=r["shape"],
                qc=r["qc_ms"],
                fp=r["fp16_ms"],
                qct=r["qc_tflops"],
                fpt=r["fp16_tflops"],
                sp=f"{float(r['speedup']):.2f}×",
                err=fmt_err(r.get("max_abs_err", "")),
            )
        )
    lines.append("")
    out.write_text("\n".join(lines), encoding="utf-8")


def write_docx(rows_w8: list[dict], rows_w4: list[dict], out: Path) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc = Document()
    h = doc.add_heading("QC-MoE vs vLLM-style FP16 MoE — Benchmark Summary (H20 GPU)", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph(
        "环境（示例）：NVIDIA H20 (sm_90)，PyTorch 2.10.0+cu128，Triton 3.6.0。"
        " 对比：QC-MoE 量化内核 vs vLLM 风格 FP16 SwiGLU MoE 参考实现。"
    )
    p.paragraph_format.space_after = Pt(8)

    cols = [
        "#",
        "Category",
        "Shape (S,H,K)",
        "QC (ms)",
        "FP16 ref (ms)",
        "QC TFLOPS",
        "FP16 TFLOPS",
        "Speedup",
        "MaxAbsErr",
    ]

    doc.add_heading("W8A16（INT8 权重 × FP16 激活）", level=2)
    t1 = doc.add_table(rows=1 + len(rows_w8), cols=len(cols))
    t1.style = "Table Grid"
    for j, c in enumerate(cols):
        t1.rows[0].cells[j].text = c
        for run in t1.rows[0].cells[j].paragraphs[0].runs:
            run.bold = True
    for i, r in enumerate(rows_w8, start=1):
        row = t1.rows[i].cells
        vals = [
            str(r["idx"]),
            r["category"],
            r["shape"],
            str(r["qc_ms"]),
            str(r["fp16_ms"]),
            str(r["qc_tflops"]),
            str(r["fp16_tflops"]),
            f"{float(r['speedup']):.2f}×",
            fmt_err(r.get("max_abs_err", "")),
        ]
        for j, v in enumerate(vals):
            row[j].text = v

    doc.add_heading("W4A16（INT4 权重 × FP16 激活）", level=2)
    t2 = doc.add_table(rows=1 + len(rows_w4), cols=len(cols))
    t2.style = "Table Grid"
    for j, c in enumerate(cols):
        t2.rows[0].cells[j].text = c
        for run in t2.rows[0].cells[j].paragraphs[0].runs:
            run.bold = True
    for i, r in enumerate(rows_w4, start=1):
        row = t2.rows[i].cells
        vals = [
            str(r["idx"]),
            r["category"],
            r["shape"],
            str(r["qc_ms"]),
            str(r["fp16_ms"]),
            str(r["qc_tflops"]),
            str(r["fp16_tflops"]),
            f"{float(r['speedup']):.2f}×",
            fmt_err(r.get("max_abs_err", "")),
        ]
        for j, v in enumerate(vals):
            row[j].text = v

    doc.save(str(out))


def main() -> None:
    env = os.environ.get("FLAGGEMS_ROOT")
    if env:
        root = Path(env)
    elif (ZH_ROOT / "qcmoe_w8a16_data.csv").is_file():
        root = ZH_ROOT
    else:
        root = DEFAULT_FLAGGEMS
    p8 = root / "qcmoe_w8a16_data.csv"
    p4 = root / "qcmoe_w4a16_data.csv"
    if not p8.is_file() or not p4.is_file():
        raise SystemExit(f"Missing CSV under {root}. Set FLAGGEMS_ROOT or copy CSVs here.")

    w8 = load_csv(p8)
    w4 = load_csv(p4)

    md = ZH_ROOT / "qc_moe_vs_vllm_fp16_benchmark_summary.md"
    dx = ZH_ROOT / "qc_moe_vs_vllm_fp16_benchmark_summary.docx"

    write_md(w8, w4, md)
    write_docx(w8, w4, dx)
    print(f"[OK] {md}")
    print(f"[OK] {dx}")


if __name__ == "__main__":
    main()
