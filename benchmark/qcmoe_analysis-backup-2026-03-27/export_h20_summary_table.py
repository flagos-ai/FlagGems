#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Export QC-MoE H20 benchmark to summary tables (Markdown + Word), styled like
the reference "Triton Benchmark Results Summary" layout.

FP8 was not measured in this MoE benchmark — FP8 columns show "—".

Usage:
    python export_h20_summary_table.py
    python export_h20_summary_table.py --w8 path/to/w8.csv --w4 path/to/w4.csv
"""

from __future__ import annotations

import argparse
import csv
import os
from typing import Any, Dict, List

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))


def load_csv(path: str) -> List[Dict[str, Any]]:
    rows = []
    with open(path, newline="", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            rows.append(row)
    return rows


def bits_label(w_nbits: int) -> str:
    if w_nbits == 8:
        return "W8A16"
    if w_nbits == 4:
        return "W4A16"
    return f"W{w_nbits}A16"


def merge_rows(w8: List[Dict], w4: List[Dict]) -> List[Dict[str, Any]]:
    """Align by shape_str + category; one row per (shape, bits)."""
    by_key = {}
    for r in w8 + w4:
        key = (r["shape_str"], int(r["w_nbits"]))
        by_key[key] = r
    # Stable order: sort by (seq*batch descending?, or by shape tuple)
    def sort_key(r):
        return (
            -int(r["seq"]) * int(r["batch"]),
            int(r["hidden"]) * int(r["inter"]),
            int(r["batch"]),
            int(r["seq"]),
            int(r["w_nbits"]),
        )

    out = list(by_key.values())
    out.sort(key=sort_key)
    return out


def enrich_row(r: Dict[str, Any]) -> Dict[str, Any]:
    batch = int(r["batch"])
    seq = int(r["seq"])
    qc_ms = float(r["qc_ms"])
    fp16_ms = float(r["fp16_ms"])
    speedup = float(r["speedup"])
    qc_tflops = float(r["qc_tflops"])
    tokens = batch * seq
    tokens_per_s = tokens / (qc_ms / 1000.0) if qc_ms > 0 else 0.0
    h, k = int(r["hidden"]), int(r["inter"])
    return {
        **r,
        "matrix_hk": f"{h}×{k}",
        "bits": bits_label(int(r["w_nbits"])),
        "tokens_total": tokens,
        "tokens_per_s": tokens_per_s,
        "speedup_fp16": speedup,
        "fp8_ms": "—",
        "speedup_fp8": "—",
        "qc_ms_4": qc_ms,
        "fp16_ms_4": fp16_ms,
        "tflops": qc_tflops,
    }


def format_ms(v: float) -> str:
    return f"{v:.4f}"


def format_speedup(x: float) -> str:
    return f"{x:.2f}x"


def format_toks(x: float) -> str:
    if x >= 1e6:
        return f"{x/1e6:.2f}M"
    if x >= 1e3:
        return f"{x/1e3:.2f}K"
    return f"{x:.0f}"


def format_tflops(x: float) -> str:
    return f"{x:.2f}"


def build_table_rows(merged: List[Dict]) -> List[Dict[str, Any]]:
    rows = [enrich_row(dict(r)) for r in merged]
    # Group by (matrix_hk, seq, batch) then bits W8 / W4 — similar to reference
    rows.sort(
        key=lambda x: (
            x["matrix_hk"],
            -x["tokens_total"],
            x["bits"],
        )
    )
    return rows


def write_markdown(rows: List[Dict[str, Any]], path: str, title: str) -> None:
    lines = [
        f"# {title}",
        "",
        "QC-MoE SwiGLU benchmark vs PyTorch FP16 reference on **NVIDIA H20**. "
        "FP8 baseline was **not** run for this workload — FP8 columns are placeholders (`—`).",
        "",
        "| Matrix (H×K) | Seq | Batch | Bits | FP16 (ms) | FP8 (ms) | Quant (ms) | Speedup vs FP16 | Speedup vs FP8 | Tokens/s | TFLOPS |",
        "|-------------:|----:|------:|:-----|----------:|:--------:|-----------:|:---------------:|:--------------:|---------:|-------:|",
    ]
    for x in rows:
        lines.append(
            "| {hk} | {seq} | {batch} | {bits} | {fp16} | {fp8} | {quant} | {sp16} | {sp8} | {toks} | {tflops} |".format(
                hk=x["matrix_hk"],
                seq=x["seq"],
                batch=x["batch"],
                bits=x["bits"],
                fp16=format_ms(x["fp16_ms_4"]),
                fp8=x["fp8_ms"],
                quant=format_ms(x["qc_ms_4"]),
                sp16=format_speedup(x["speedup_fp16"]),
                sp8=x["speedup_fp8"],
                toks=format_toks(x["tokens_per_s"]),
                tflops=format_tflops(x["tflops"]),
            )
        )
    lines.extend(
        [
            "",
            "---",
            "",
            "## Column notes",
            "",
            "- **Matrix (H×K)**: hidden × intermediate for the MoE FFN projection footprint.",
            "- **Seq / Batch**: from YAML shape `(B, S, H, K)`; effective token count = B×S.",
            "- **Quant (ms)**: FlagGems QC-MoE Triton kernel (W4A16 / W8A16).",
            "- **TFLOPS**: from benchmark FLOPs / latency (same definition as CSV).",
            "",
        ]
    )
    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))


def write_docx(rows: List[Dict[str, Any]], path: str, title: str) -> None:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    def set_cell_shading(cell, fill_hex: str) -> None:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill_hex)
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_pr.append(shd)

    doc = Document()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph(
        "QC-MoE SwiGLU vs PyTorch FP16 reference on NVIDIA H20. "
        "FP8 baseline not measured — FP8 cells show em dash."
    )

    headers = [
        "Matrix (H×K)",
        "Seq",
        "Batch",
        "Bits",
        "FP16 (ms)",
        "FP8 (ms)",
        "Quant (ms)",
        "Speedup vs FP16",
        "Speedup vs FP8",
        "Tokens/s",
        "TFLOPS",
    ]
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        c = hdr_cells[i]
        c.text = ""
        hp = c.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr = hp.add_run(h)
        hr.bold = True
        hr.font.color.rgb = RGBColor(255, 255, 255)
        hr.font.size = Pt(9)
        set_cell_shading(c, "2E5090")

    for ri, x in enumerate(rows, start=1):
        cells = table.rows[ri].cells
        vals = [
            x["matrix_hk"],
            str(x["seq"]),
            str(x["batch"]),
            x["bits"],
            format_ms(x["fp16_ms_4"]),
            "—",
            format_ms(x["qc_ms_4"]),
            format_speedup(x["speedup_fp16"]),
            "—",
            format_toks(x["tokens_per_s"]),
            format_tflops(x["tflops"]),
        ]
        stripe = "F2F2F2" if ri % 2 == 0 else "FFFFFF"
        for ci, v in enumerate(vals):
            c = cells[ci]
            c.text = str(v)
            set_cell_shading(c, stripe)
            for para in c.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in para.runs:
                    r.font.size = Pt(8)

    doc.save(path)


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument(
        "--w8",
        default=os.path.join(SCRIPT_DIR, "qcmoe_complete_w8a16_data.csv"),
    )
    ap.add_argument(
        "--w4",
        default=os.path.join(SCRIPT_DIR, "qcmoe_complete_w4a16_data.csv"),
    )
    ap.add_argument(
        "--md",
        default=os.path.join(SCRIPT_DIR, "qcmoe_h20_summary_table.md"),
    )
    ap.add_argument(
        "--docx",
        default=os.path.join(SCRIPT_DIR, "qcmoe_h20_summary_table.docx"),
    )
    args = ap.parse_args()

    w8 = load_csv(args.w8)
    w4 = load_csv(args.w4)
    merged = merge_rows(w8, w4)
    rows = build_table_rows(merged)

    title = "FlagGems QC-MoE Benchmark Summary (H20 GPU)"
    write_markdown(rows, args.md, title)
    write_docx(rows, args.docx, title)
    print(f"[OK] Markdown: {args.md}")
    print(f"[OK] Word:     {args.docx}")
    print(f"     Rows:    {len(rows)}")


if __name__ == "__main__":
    main()
