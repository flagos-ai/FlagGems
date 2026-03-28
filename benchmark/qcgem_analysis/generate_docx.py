#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate benchmark_analysis.docx (Word version of benchmark_analysis.md)
and qcgem_complete_report.docx for FlagGems QC-GEM analysis.
"""

from __future__ import annotations

import csv
import os
from pathlib import Path

SCRIPT_DIR = Path(__file__).resolve().parent
CSV_W8 = SCRIPT_DIR / "qcgem_w8a16_data.csv"
CSV_W4 = SCRIPT_DIR / "qcgem_w4a16_data.csv"


def load_csv(path: Path):
    with open(path, encoding="utf-8", newline="") as f:
        return list(csv.DictReader(f))


def fmt_err(v: str) -> str:
    v = (v or "").strip().lower()
    if v in ("nan", "", "none"):
        return "—"
    return v


def write_benchmark_analysis_docx(out: Path) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    def set_shading(cell, fill_hex: str) -> None:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill_hex)
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_pr.append(shd)

    doc = Document()
    doc.add_heading("QC-GEM Benchmark Analysis Report", level=1)

    doc.add_paragraph(
        "Author: QC-GEM Analysis Script  |  Date: 2026-03-27  |  "
        "Test Environment: Qwen3.5-397B-A17B MoE Shapes"
    )
    doc.add_paragraph(
        "This report presents the performance analysis of QC-GEM kernel "
        "implementation compared to PyTorch baseline for quantization-aware GEMM operations. "
        "Two quantization configurations were evaluated: W4A16 and W8A16."
    )

    # ── Key Findings ──────────────────────────────────────────────────────────
    doc.add_heading("1. Executive Summary", level=2)
    doc.add_heading("Key Findings", level=3)

    headers_kf = ["Metric", "w4A16", "w8A16", "Difference"]
    kf_data = [
        ["Average Speedup", "0.895x", "0.915x", "-0.020"],
        ["Max Speedup", "1.041x", "1.026x", "+0.015"],
        ["Min Speedup", "0.603x", "0.585x", "+0.018"],
        ["Average TFLOPS", "59.27", "55.15", "+4.11"],
        ["Max TFLOPS", "93.84", "93.84", "0.00"],
    ]
    t = doc.add_table(rows=1 + len(kf_data), cols=4)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers_kf):
        c = t.rows[0].cells[j]
        c.text = h
        set_shading(c, "2E5090")
        for r in c.paragraphs[0].runs:
            r.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.size = Pt(9)
    for ri, row_data in enumerate(kf_data, 1):
        stripe = "F2F2F2" if ri % 2 == 0 else "FFFFFF"
        for ci, v in enumerate(row_data):
            c = t.rows[ri].cells[ci]
            c.text = v
            set_shading(c, stripe)
            for para in c.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in para.runs:
                    r.font.size = Pt(9)

    doc.add_paragraph(
        "Conclusion: w8A16 achieves slightly higher average speedup (+2%) compared to w4A16, "
        "but w4A16 achieves higher peak performance on large shapes."
    )

    # ── Performance by Category ─────────────────────────────────────────────
    doc.add_heading("2. Performance Analysis", level=2)

    headers_cat = ["Category", "W8A16 Speedup Range", "W4A16 Speedup Range", "Winner"]
    cat_data = [
        ["Down projection M≥4096", "1.01x ~ 1.03x", "1.01x ~ 1.04x", "W4A16"],
        ["Down projection M≤2048", "0.75x ~ 0.96x", "0.64x ~ 0.84x", "W8A16"],
        ["Up projection M≥4096", "1.00x ~ 1.01x", "1.00x ~ 1.02x", "W4A16"],
        ["Up projection M≤2048", "0.96x ~ 0.96x", "0.82x ~ 0.82x", "W8A16"],
        ["Gate projection", "0.98x ~ 1.01x", "1.00x ~ 1.01x", "W4A16"],
        ["Router (N=128)", "0.59x ~ 0.67x", "0.60x ~ 0.73x", "W4A16"],
    ]
    t2 = doc.add_table(rows=1 + len(cat_data), cols=4)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers_cat):
        c = t2.rows[0].cells[j]
        c.text = h
        set_shading(c, "2E5090")
        for r in c.paragraphs[0].runs:
            r.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.size = Pt(9)
    for ri, row_data in enumerate(cat_data, 1):
        stripe = "F2F2F2" if ri % 2 == 0 else "FFFFFF"
        for ci, v in enumerate(row_data):
            c = t2.rows[ri].cells[ci]
            c.text = v
            set_shading(c, stripe)
            for para in c.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in para.runs:
                    r.font.size = Pt(9)

    # ── Recommendations ────────────────────────────────────────────────────────
    doc.add_heading("3. Mode Selection Recommendations", level=2)
    doc.add_heading("When to Use w4A16", level=3)
    for item in [
        "Large batch sizes (M ≥ 4096) with FFN layers",
        "Maximum throughput requirements",
        "Memory bandwidth limited scenarios (higher compression ratio)",
        "FFN layer computations with N ≥ 1024",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("When to Use w8A16", level=3)
    for item in [
        "Small batch sizes (M < 4096)",
        "Router layer computations",
        "Cases requiring better small-shape performance",
        "When weight quantization precision matters for accuracy",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("When to Fall Back to PyTorch", level=3)
    for item in [
        "Router layer computations (N = 128)",
        "Very small shapes (M < 512) where overhead dominates",
        "Low-latency requirements with small shapes",
        "Cases where quantization overhead exceeds compute savings",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.save(str(out))


def write_complete_report_docx(out: Path) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    def set_shading(cell, fill_hex: str) -> None:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill_hex)
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_pr.append(shd)

    doc = Document()
    doc.add_heading("QC-GEM Benchmark: 完整测试报告", level=1)
    doc.add_paragraph(
        "FlagGems Triton Kernel vs PyTorch FP16 GEMM  |  "
        "日期: 2026-03-27  |  GPU: NVIDIA H20  |  模型: Qwen3.5-397B-A17B"
    )

    # ── W8A16 table ─────────────────────────────────────────────────────────
    doc.add_heading("W8A16 数据表", level=2)
    w8 = load_csv(CSV_W8)
    cols = ["#", "Category", "Shape", "QC W8A16 (ms)", "FP16 ref (ms)", "QC TFLOPS", "FP16 TFLOPS", "Speedup"]
    t = doc.add_table(rows=1 + len(w8), cols=len(cols))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(cols):
        c = t.rows[0].cells[j]
        c.text = h
        set_shading(c, "2E5090")
        for r in c.paragraphs[0].runs:
            r.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.size = Pt(8)
    for ri, r in enumerate(w8, 1):
        stripe = "F2F2F2" if ri % 2 == 0 else "FFFFFF"
        vals = [str(r["idx"]), r["category"], r["shape"],
                r["qc_ms"], r["fp16_ms"], r["qc_tflops"],
                r["fp16_tflops"], f"{float(r['speedup']):.3f}×"]
        for ci, v in enumerate(vals):
            c = t.rows[ri].cells[ci]
            c.text = v
            set_shading(c, stripe)
            for para in c.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r2 in para.runs:
                    r2.font.size = Pt(8)

    # ── W4A16 table ─────────────────────────────────────────────────────────
    doc.add_heading("W4A16 数据表", level=2)
    w4 = load_csv(CSV_W4)
    t2 = doc.add_table(rows=1 + len(w4), cols=len(cols))
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(cols):
        c = t2.rows[0].cells[j]
        c.text = h
        set_shading(c, "2E5090")
        for r in c.paragraphs[0].runs:
            r.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.size = Pt(8)
    for ri, r in enumerate(w4, 1):
        stripe = "F2F2F2" if ri % 2 == 0 else "FFFFFF"
        vals = [str(r["idx"]), r["category"], r["shape"],
                r["qc_ms"], r["fp16_ms"], r["qc_tflops"],
                r["fp16_tflops"], f"{float(r['speedup']):.3f}×"]
        for ci, v in enumerate(vals):
            c = t2.rows[ri].cells[ci]
            c.text = v
            set_shading(c, stripe)
            for para in c.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r2 in para.runs:
                    r2.font.size = Pt(8)

    doc.save(str(out))


def main() -> None:
    write_benchmark_analysis_docx(SCRIPT_DIR / "benchmark_analysis.docx")
    print(f"[OK] benchmark_analysis.docx")
    write_complete_report_docx(SCRIPT_DIR / "qcgem_complete_report.docx")
    print(f"[OK] qcgem_complete_report.docx")


if __name__ == "__main__":
    main()
