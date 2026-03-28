#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Build QC-GEM (FlagGems qcgem) summary tables in the same layout as
benchmark/qcmoe_analysis/qcmoe_h20_summary_table.md.

CSV columns: m,n,k,torch_latency_ms,gems_latency_ms,speedup,tflops,...

FP8 baseline was not run — FP8 columns show em dash (—).

Usage:
    python3 export_qcgem_summary_table.py
    python3 export_qcgem_summary_table.py --w4 path.csv --w8 path.csv
"""

from __future__ import annotations

import argparse
import csv
import os
from typing import Any, Dict, List, Tuple

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))


def load_qcgem_csv(path: str) -> List[Dict[str, Any]]:
    rows = []
    with open(path, newline="", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            rows.append(row)
    return rows


def matrix_hk(n: int, k: int) -> str:
    """Match Qwen MoE YAML naming (same as QC-MoE summary table)."""
    if n == 1024 and k == 3584:
        return "1024×3584"
    if n == 3584 and k == 1024:
        return "3584×1024"
    if n == 7168 and k == 1024:
        return "1024×7168"
    if n == 128 and k == 1024:
        return "128×1024"
    return f"{n}×{k}"


def row_from_csv(r: Dict[str, Any], bits: str, key: Tuple[int, int, int, str]) -> Dict[str, Any]:
    m = int(r["m"])
    n = int(r["n"])
    k = int(r["k"])
    fp16_ms = float(r["torch_latency_ms"])
    quant_ms = float(r["gems_latency_ms"])
    speedup = float(r["speedup"])
    tflops = float(r["tflops"])
    seq = m
    batch = 1
    tokens = m
    tokens_per_s = tokens / (quant_ms / 1000.0) if quant_ms > 0 else 0.0
    return {
        "key": key,
        "matrix_hk": matrix_hk(n, k),
        "seq": seq,
        "batch": batch,
        "bits": bits,
        "fp16_ms": fp16_ms,
        "quant_ms": quant_ms,
        "speedup_fp16": speedup,
        "tokens_per_s": tokens_per_s,
        "tflops": tflops,
        "fp8_ms": "—",
        "speedup_fp8": "—",
    }


def merge_w4_w8(
    w4_rows: List[Dict[str, Any]],
    w8_rows: List[Dict[str, Any]],
) -> List[Dict[str, Any]]:
    by_key: Dict[Tuple[int, int, int, str], Dict[str, Any]] = {}
    for r in w8_rows:
        m, n, k = int(r["m"]), int(r["n"]), int(r["k"])
        key = (m, n, k, "W8A16")
        by_key[key] = row_from_csv(r, "W8A16", key)
    for r in w4_rows:
        m, n, k = int(r["m"]), int(r["n"]), int(r["k"])
        key = (m, n, k, "W4A16")
        by_key[key] = row_from_csv(r, "W4A16", key)
    out = list(by_key.values())

    def sort_key(x: Dict[str, Any]) -> Any:
        m = x["key"][0]
        n = x["key"][1]
        k = x["key"][2]
        bits_order = 0 if x["bits"] == "W8A16" else 1
        return (
            x["matrix_hk"],
            -m * x["batch"],
            bits_order,
            -m,
            n,
            k,
        )

    out.sort(key=sort_key)
    return out


def format_ms(v: float) -> str:
    return f"{v:.4f}"


def format_speedup(x: float) -> str:
    return f"{x:.2f}x"


def format_toks(x: float) -> str:
    if x >= 1e6:
        return f"{x/1e6:.2f}M"
    if x >= 1e3:
        return f"{x/1e3:.2f}K"
    return f"{x:.0f}"


def format_tflops(x: float) -> str:
    return f"{x:.2f}"


def write_markdown(rows: List[Dict[str, Any]], path: str, title: str) -> None:
    lines = [
        f"# {title}",
        "",
        "QC-GEM quantized GEMM benchmark vs PyTorch FP16 reference on **NVIDIA H20** "
        "(same table layout as QC-MoE summary). Shapes follow Qwen3.5 MoE YAML `(M, N, K)` "
        "with **M = batch×seq** tokens. FP8 baseline was **not** run — FP8 columns are placeholders (`—`).",
        "",
        "| Matrix (H×K) | Seq | Batch | Bits | FP16 (ms) | FP8 (ms) | Quant (ms) | Speedup vs FP16 | Speedup vs FP8 | Tokens/s | TFLOPS |",
        "|-------------:|----:|------:|:-----|----------:|:--------:|-----------:|:---------------:|:--------------:|---------:|-------:|",
    ]
    for x in rows:
        lines.append(
            "| {hk} | {seq} | {batch} | {bits} | {fp16} | {fp8} | {quant} | {sp16} | {sp8} | {toks} | {tflops} |".format(
                hk=x["matrix_hk"],
                seq=x["seq"],
                batch=x["batch"],
                bits=x["bits"],
                fp16=format_ms(x["fp16_ms"]),
                fp8=x["fp8_ms"],
                quant=format_ms(x["quant_ms"]),
                sp16=format_speedup(x["speedup_fp16"]),
                sp8=x["speedup_fp8"],
                toks=format_toks(x["tokens_per_s"]),
                tflops=format_tflops(x["tflops"]),
            )
        )
    lines.extend(
        [
            "",
            "---",
            "",
            "## Column notes",
            "",
            "- **Matrix (H×K)**: MoE projection footprint (`N×K` from YAML, same labels as QC-MoE: e.g. down `1024×3584`, up `3584×1024`).",
            "- **Seq / Batch**: This benchmark flattens tokens into **M**; we show **Seq = M**, **Batch = 1** (effective tokens = M).",
            "- **FP16 (ms)**: PyTorch FP16 reference (`torch_latency_ms` in CSV).",
            "- **Quant (ms)**: FlagGems QC-GEM Triton path (`gems_latency_ms`).",
            "- **Speedup vs FP16**: FP16 / Quant.",
            "- **TFLOPS**: From benchmark CSV (quant path), same definition as plotting scripts.",
            "",
        ]
    )
    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))


def write_docx(rows: List[Dict[str, Any]], path: str, title: str) -> None:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    def set_cell_shading(cell, fill_hex: str) -> None:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill_hex)
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_pr.append(shd)

    doc = Document()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph(
        "QC-GEM vs PyTorch FP16 on H20. Shapes (M,N,K) from Qwen3.5 MoE YAML; Seq=M, Batch=1. "
        "FP8 not measured — FP8 cells show em dash."
    )

    headers = [
        "Matrix (H×K)",
        "Seq",
        "Batch",
        "Bits",
        "FP16 (ms)",
        "FP8 (ms)",
        "Quant (ms)",
        "Speedup vs FP16",
        "Speedup vs FP8",
        "Tokens/s",
        "TFLOPS",
    ]
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        c = hdr_cells[i]
        c.text = ""
        hp = c.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr = hp.add_run(h)
        hr.bold = True
        hr.font.color.rgb = RGBColor(255, 255, 255)
        hr.font.size = Pt(9)
        set_cell_shading(c, "2E5090")

    for ri, x in enumerate(rows, start=1):
        cells = table.rows[ri].cells
        vals = [
            x["matrix_hk"],
            str(x["seq"]),
            str(x["batch"]),
            x["bits"],
            format_ms(x["fp16_ms"]),
            "—",
            format_ms(x["quant_ms"]),
            format_speedup(x["speedup_fp16"]),
            "—",
            format_toks(x["tokens_per_s"]),
            format_tflops(x["tflops"]),
        ]
        stripe = "F2F2F2" if ri % 2 == 0 else "FFFFFF"
        for ci, v in enumerate(vals):
            c = cells[ci]
            c.text = str(v)
            set_cell_shading(c, stripe)
            for para in c.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in para.runs:
                    r.font.size = Pt(8)

    doc.save(path)


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument(
        "--w4",
        default=os.path.join(SCRIPT_DIR, "benchmark_w4a16_fp16.csv"),
    )
    ap.add_argument(
        "--w8",
        default=os.path.join(SCRIPT_DIR, "benchmark_w8a16_fp16.csv"),
    )
    ap.add_argument(
        "--md",
        default=os.path.join(SCRIPT_DIR, "qcgem_summary_table.md"),
    )
    ap.add_argument(
        "--docx",
        default=os.path.join(SCRIPT_DIR, "qcgem_summary_table.docx"),
    )
    args = ap.parse_args()

    w4 = load_qcgem_csv(args.w4)
    w8 = load_qcgem_csv(args.w8)
    rows = merge_w4_w8(w4, w8)

    title = "FlagGems QC-GEM Benchmark Summary (H20 GPU)"
    write_markdown(rows, args.md, title)
    write_docx(rows, args.docx, title)
    print(f"[OK] Markdown: {args.md}")
    print(f"[OK] Word:     {args.docx}")
    print(f"     Rows:    {len(rows)}")


if __name__ == "__main__":
    main()
